
import streamlit as st
import pandas as pd
from openai import OpenAI
from google.cloud import bigquery
import PyPDF2
import docx
import numpy as np
from numpy.linalg import norm
import re
import json

# ============================================================
# CONFIGURATION
# ============================================================
PROJECT_ID = "ba882-team4-474802"
DATASET_ID = "ba882_jobs"
EMBED_MODEL = "text-embedding-3-small"
MODEL_NAME = "gpt-4o-mini"

st.set_page_config(page_title="Career Assistant Suite", layout="wide")

# ============================================================
# AUTH — OPENAI
# ============================================================
if "OPEN_AI_API" not in st.secrets:
    st.error("❌ Missing OPEN_AI_API in secrets.toml")
else:
    client = OpenAI(api_key=st.secrets["OPEN_AI_API"])

# ============================================================
# BIGQUERY CLIENT
# ============================================================
client_bq = bigquery.Client(project=PROJECT_ID)

# ============================================================
# HELPER — Extract Resume Text
# ============================================================
def extract_text_from_file(uploaded_file):
    text = ""
    try:
        if uploaded_file.name.endswith(".pdf"):
            reader = PyPDF2.PdfReader(uploaded_file)
            for page in reader.pages:
                extracted = page.extract_text()
                if extracted:
                    text += extracted + "\n"

        elif uploaded_file.name.endswith(".docx"):
            doc = docx.Document(uploaded_file)
            for para in doc.paragraphs:
                text += para.text + "\n"

        else:
            text = uploaded_file.read().decode("utf-8", errors="ignore")

    except Exception as e:
        st.error(f"Error reading file: {e}")

    return text

# ============================================================
# SAFE JSON LIST EXTRACTOR
# ============================================================
def safe_extract_keywords(text):
    match = re.search(r"\[[^\]]*\]", text)
    if not match:
        return []
    try:
        return json.loads(match.group(0))
    except:
        return []

# ============================================================
# LLM Keyword Extraction
# ============================================================
def extract_keywords_with_llm(resume_text, max_keywords=10):
    prompt = f"""
    Extract the {max_keywords} most relevant job-related keywords from the resume.

    OUTPUT RULES:
    - You MUST output ONLY a JSON array.
    - Example: ["python", "sql", "finance"]

    Resume:
    {resume_text}
    """

    response = client.chat.completions.create(
        model=MODEL_NAME,
        messages=[
            {"role": "system", "content": "Output only a JSON array of keywords."},
            {"role": "user", "content": prompt}
        ]
    )

    raw = response.choices[0].message.content.strip()

    try:
        arr = json.loads(raw)
        if isinstance(arr, list):
            return [k.lower() for k in arr]
    except:
        pass

    arr = safe_extract_keywords(raw)
    return [k.lower() for k in arr] if isinstance(arr, list) else []


# ============================================================
# Embedding helpers
# ============================================================
def load_job_embeddings_filtered(selected_state):
    if selected_state == "All States":
        query = f"""
            SELECT job_id, embedding_vector
            FROM {PROJECT_ID}.{DATASET_ID}.job_embeddings
        """
    else:
        query = f"""
            SELECT e.job_id, e.embedding_vector
            FROM {PROJECT_ID}.{DATASET_ID}.job_embeddings e
            JOIN {PROJECT_ID}.{DATASET_ID}.jobs j
            ON e.job_id = j.job_id
            WHERE j.state = '{selected_state}'
        """
    return client_bq.query(query).to_dataframe()


# ============================================================
# fetch_job_details (includes sponsor)
# ============================================================
def fetch_job_details(job_ids):
    formatted = ",".join([f"'{jid}'" for jid in job_ids])

    query = f"""
        SELECT
            j.job_id,
            j.title,
            j.company_name,
            j.company_id,
            j.city,
            j.state,
            j.description,
            j.salary_min,
            j.salary_max,
            j.redirected_url,
            sc.is_sponsor
        FROM {PROJECT_ID}.{DATASET_ID}.jobs j
        LEFT JOIN {PROJECT_ID}.{DATASET_ID}.sponsor_companies sc
        ON j.company_id = sc.company_id
        WHERE j.job_id IN ({formatted})
    """

    return client_bq.query(query).to_dataframe()

def embed_text(text):
    emb = client.embeddings.create(model=EMBED_MODEL, input=text)
    return np.array(emb.data[0].embedding)

def cosine_similarity(a, b):
    return np.dot(a, b) / (norm(a) * norm(b))


# ============================================================
# RAG Context (from old version)
# ============================================================
def retrieve_rag_context(query_text, top_k=3):
    tokens = [w.lower() for w in query_text.split()]
    keywords = [w for w in tokens if len(w) > 3]

    if not keywords:
        return "No meaningful terms found."

    keyword = keywords[0]

    query = f"""
        SELECT title, company_name, description
        FROM {PROJECT_ID}.{DATASET_ID}.jobs
        WHERE LOWER(description) LIKE LOWER('%{keyword}%')
           OR LOWER(title) LIKE LOWER('%{keyword}%')
        LIMIT {int(top_k)}
    """

    df = client_bq.query(query).to_dataframe()

    if df.empty:
        return "No similar job descriptions found."

    context_str = ""
    for _, row in df.iterrows():
        desc = row.get("description", "")
        if not isinstance(desc, str):
            desc = ""
        context_str += f"Title: {row['title']} | Company: {row['company_name']}\n"
        context_str += f"Description Snippet: {desc[:700]}...\n\n"

    return context_str


# ============================================================
# UI: Two Tabs
# ============================================================
tab1, tab2 = st.tabs(["📄 Resume Matcher", "🤖 Hiring Assistant"])


# ============================================================
# TAB 1 (New Resume Matcher – DO NOT MODIFY)
# ============================================================
with tab1:
    st.title("📄 AI-Powered Resume Matcher")

    states_df = client_bq.query(
        f"SELECT DISTINCT state FROM {PROJECT_ID}.{DATASET_ID}.jobs WHERE state IS NOT NULL"
    ).to_dataframe()

    selected_state = st.selectbox("Filter by State (Optional)",
                                  ["All States"] + sorted(states_df["state"].unique()))

    uploaded_file = st.file_uploader("Upload Resume", type=["pdf", "docx", "txt"])

    if uploaded_file:
        resume_text = extract_text_from_file(uploaded_file)
        st.success("Resume parsed successfully!")
        st.text_area("Extracted Resume Text", resume_text, height=200)

        if st.button("🔍 Find Best Job Matches"):
            with st.spinner("Processing..."):

                keywords = extract_keywords_with_llm(resume_text)
                st.subheader("🔑 Extracted Keywords")
                st.write(", ".join(keywords))

                resume_emb = embed_text(resume_text)
                job_emb_df = load_job_embeddings_filtered(selected_state)

                sims = []
                for _, r in job_emb_df.iterrows():
                    vec = np.array(r["embedding_vector"])
                    sims.append((r["job_id"], cosine_similarity(resume_emb, vec)))

                sims = sorted(sims, key=lambda x: x[1], reverse=True)

                top_ids = [jid for jid, _ in sims[:20]]
                job_details = fetch_job_details(top_ids)

                job_details["similarity"] = job_details["job_id"].map(dict(sims))

                job_details["unique_key"] = (
                    job_details["title"].str.lower() + " — " + job_details["company_name"].str.lower()
                )

                top3 = (
                    job_details.drop_duplicates("unique_key")
                    .sort_values("similarity", ascending=False)
                    .head(3)
                )

            st.subheader("🏆 Top 3 Job Recommendations")

            for _, row in top3.iterrows():
                with st.expander(f"{row['title']} — {row['company_name']}"):

                    st.markdown("### 📍 Location")
                    st.write(f"{row['city']}, {row['state']}")

                    st.markdown("### 📝 Description")
                    st.write(row["description"][:700] + "...")

                    st.markdown("### 💰 Salary Range")
                    st.write(f"{row['salary_min']} — {row['salary_max']}")

                    st.markdown("### 🔗 Apply Link")
                    st.write(f"[View Job Posting]({row['redirected_url']})")

                    st.markdown("### 🛂 Sponsorship Eligibility")

                    if row["is_sponsor"] is True:
                        st.markdown("<span style='font-size:22px;'>✅ Yes — H1B Sponsorship Available</span>",
                                    unsafe_allow_html=True)
                    elif row["is_sponsor"] is False:
                        st.markdown("<span style='font-size:22px;'>❌ No — Sponsorship Not Offered</span>",
                                    unsafe_allow_html=True)
                    else:
                        st.markdown("<span style='font-size:22px;'>❓ Unknown — No Sponsorship Data</span>",
                                    unsafe_allow_html=True)


# ============================================================
# TAB 2 — Hiring Assistant (UNCHANGED)
# ============================================================
with tab2:
    st.title("🤖 Hiring Assistant (RAG Chatbot)")
    st.caption("Ask me to draft job descriptions or hiring messages using real job market data.")

    if "messages" not in st.session_state:
        st.session_state.messages = [{
            "role": "assistant",
            "content": "Hello! What role are you hiring for? (e.g., 'Senior Data Analyst in NYC')"
        }]

    for message in st.session_state.messages:
        with st.chat_message(message["role"]):
            st.markdown(message["content"])

    if prompt := st.chat_input("Describe the role..."):
        st.session_state.messages.append({"role": "user", "content": prompt})

        with st.chat_message("user"):
            st.markdown(prompt)

        with st.chat_message("assistant"):
            with st.spinner("Retrieving job market data..."):
                context_data = retrieve_rag_context(prompt)

                system_prompt = f"""
                You are an expert HR assistant.
                Use the REAL, RECENT job postings below to generate insights.

                --- MARKET CONTEXT ---
                {context_data}
                ----------------------

                USER REQUEST: {prompt}

                Your job: Write a polished, professional job description or hiring message.
                """

                response = client.chat.completions.create(
                    model=MODEL_NAME,
                    messages=[{"role": "system", "content": system_prompt}]
                )

                output = response.choices[0].message.content
                st.markdown(output)

        st.session_state.messages.append({"role": "assistant", "content": output})


# import streamlit as st
# import pandas as pd
# from openai import OpenAI
# from google.cloud import bigquery
# import PyPDF2
# import docx
# import numpy as np
# from numpy.linalg import norm
# import re
# import json

# # ============================================================
# # CONFIGURATION
# # ============================================================
# PROJECT_ID = "ba882-team4-474802"
# DATASET_ID = "ba882_jobs"
# EMBED_MODEL = "text-embedding-3-small"
# MODEL_NAME = "gpt-4o-mini"

# st.set_page_config(page_title="Career Assistant Suite", layout="wide")

# # ============================================================
# # AUTH — OPENAI
# # ============================================================
# if "OPEN_AI_API" not in st.secrets:
#     st.error("❌ Missing OPEN_AI_API in secrets.toml")
# else:
#     client = OpenAI(api_key=st.secrets["OPEN_AI_API"])

# # ============================================================
# # BIGQUERY CLIENT
# # ============================================================
# client_bq = bigquery.Client(project=PROJECT_ID)

# # ============================================================
# # HELPER — Extract Resume Text
# # ============================================================
# def extract_text_from_file(uploaded_file):
#     text = ""
#     try:
#         if uploaded_file.name.endswith(".pdf"):
#             reader = PyPDF2.PdfReader(uploaded_file)
#             for page in reader.pages:
#                 extracted = page.extract_text()
#                 if extracted:
#                     text += extracted + "\n"

#         elif uploaded_file.name.endswith(".docx"):
#             doc = docx.Document(uploaded_file)
#             for para in doc.paragraphs:
#                 text += para.text + "\n"

#         else:
#             text = uploaded_file.read().decode("utf-8", errors="ignore")

#     except Exception as e:
#         st.error(f"Error reading file: {e}")

#     return text

# # ============================================================
# # SAFE JSON LIST EXTRACTOR
# # ============================================================
# def safe_extract_keywords(text):
#     match = re.search(r"\[[^\]]*\]", text)
#     if not match:
#         return []
#     try:
#         return json.loads(match.group(0))
#     except:
#         return []

# # ============================================================
# # LLM Keyword Extraction
# # ============================================================
# def extract_keywords_with_llm(resume_text, max_keywords=10):
#     prompt = f"""
#     Extract the {max_keywords} most relevant job-related keywords from the resume.

#     OUTPUT RULES:
#     - You MUST output ONLY a JSON array.
#     - Example: ["python", "sql", "finance"]

#     Resume:
#     {resume_text}
#     """

#     response = client.chat.completions.create(
#         model=MODEL_NAME,
#         messages=[
#             {"role": "system", "content": "Output only a JSON array of keywords."},
#             {"role": "user", "content": prompt}
#         ]
#     )

#     raw = response.choices[0].message.content.strip()

#     # direct JSON parse
#     try:
#         arr = json.loads(raw)
#         if isinstance(arr, list):
#             return [k.lower() for k in arr]
#     except:
#         pass

#     # fallback safe parse
#     arr = safe_extract_keywords(raw)
#     return [k.lower() for k in arr] if isinstance(arr, list) else []


# # ============================================================
# # Embedding helpers
# # ============================================================
# def load_job_embeddings_filtered(selected_state):
#     if selected_state == "All States":
#         query = f"""
#             SELECT job_id, embedding_vector
#             FROM {PROJECT_ID}.{DATASET_ID}.job_embeddings
#         """
#     else:
#         query = f"""
#             SELECT e.job_id, e.embedding_vector
#             FROM {PROJECT_ID}.{DATASET_ID}.job_embeddings e
#             JOIN {PROJECT_ID}.{DATASET_ID}.jobs j
#             ON e.job_id = j.job_id
#             WHERE j.state = '{selected_state}'
#         """
#     return client_bq.query(query).to_dataframe()


# # ============================================================
# # UPDATED — fetch_job_details (includes sponsor info)
# # ============================================================
# def fetch_job_details(job_ids):
#     formatted = ",".join([f"'{jid}'" for jid in job_ids])

#     query = f"""
#         SELECT
#             j.job_id,
#             j.title,
#             j.company_name,
#             j.company_id,
#             j.city,
#             j.state,
#             j.description,
#             j.salary_min,
#             j.salary_max,
#             j.redirected_url,
#             sc.is_sponsor
#         FROM {PROJECT_ID}.{DATASET_ID}.jobs j
#         LEFT JOIN {PROJECT_ID}.{DATASET_ID}.sponsor_companies sc
#         ON j.company_id = sc.company_id
#         WHERE j.job_id IN ({formatted})
#     """

#     return client_bq.query(query).to_dataframe()


# def embed_text(text):
#     emb = client.embeddings.create(model=EMBED_MODEL, input=text)
#     return np.array(emb.data[0].embedding)


# def cosine_similarity(a, b):
#     return np.dot(a, b) / (norm(a) * norm(b))


# # ============================================================
# # RAG Context
# # ============================================================
# def retrieve_rag_context(query_text, top_k=3):
#     tokens = [w.lower() for w in query_text.split()]
#     keywords = [w for w in tokens if len(w) > 3]

#     if not keywords:
#         return "No meaningful terms found."

#     kw = keywords[0]

#     query = f"""
#         SELECT title, company_name, description
#         FROM {PROJECT_ID}.{DATASET_ID}.jobs
#         WHERE LOWER(description) LIKE '%{kw}%'
#            OR LOWER(title) LIKE '%{kw}%'
#         LIMIT {top_k}
#     """

#     df = client_bq.query(query).to_dataframe()
#     if df.empty:
#         return "No similar jobs found."

#     out = ""
#     for _, row in df.iterrows():
#         out += f"Title: {row['title']} | Company: {row['company_name']}\n"
#         out += f"Description Snippet: {row['description'][:700]}...\n\n"

#     return out


# # ============================================================
# # UI TABS
# # ============================================================
# tab1, tab2 = st.tabs(["📄 Resume Matcher", "🤖 Hiring Assistant"])

# # ============================================================
# # TAB 1 — Resume Matcher
# # ============================================================
# with tab1:
#     st.title("📄 AI-Powered Resume Matcher")

#     states_df = client_bq.query(
#         f"SELECT DISTINCT state FROM {PROJECT_ID}.{DATASET_ID}.jobs WHERE state IS NOT NULL"
#     ).to_dataframe()

#     selected_state = st.selectbox("Filter by State (Optional)",
#                                   ["All States"] + sorted(states_df["state"].unique()))

#     uploaded_file = st.file_uploader("Upload Resume", type=["pdf", "docx", "txt"])

#     if uploaded_file:
#         resume_text = extract_text_from_file(uploaded_file)
#         st.success("Resume parsed successfully!")
#         st.text_area("Extracted Resume Text", resume_text, height=200)

#         if st.button("🔍 Find Best Job Matches"):
#             with st.spinner("Processing..."):

#                 keywords = extract_keywords_with_llm(resume_text)
#                 st.subheader("🔑 Extracted Keywords")
#                 st.write(", ".join(keywords))

#                 resume_emb = embed_text(resume_text)
#                 job_emb_df = load_job_embeddings_filtered(selected_state)

#                 sims = []
#                 for _, r in job_emb_df.iterrows():
#                     vec = np.array(r["embedding_vector"])
#                     sims.append((r["job_id"], cosine_similarity(resume_emb, vec)))

#                 sims = sorted(sims, key=lambda x: x[1], reverse=True)

#                 top_ids = [jid for jid, _ in sims[:20]]
#                 job_details = fetch_job_details(top_ids)

#                 job_details["similarity"] = job_details["job_id"].map(dict(sims))

#                 job_details["unique_key"] = (
#                     job_details["title"].str.lower() + " — " + job_details["company_name"].str.lower()
#                 )

#                 top3 = (
#                     job_details.drop_duplicates("unique_key")
#                     .sort_values("similarity", ascending=False)
#                     .head(3)
#                 )

#             st.subheader("🏆 Top 3 Job Recommendations")

#             for _, row in top3.iterrows():
#                 with st.expander(f"{row['title']} — {row['company_name']}"):

#                     st.markdown("### 📍 Location")
#                     st.write(f"{row['city']}, {row['state']}")

#                     st.markdown("### 📝 Description")
#                     st.write(row["description"][:700] + "...")

#                     st.markdown("### 💰 Salary Range")
#                     st.write(f"{row['salary_min']} — {row['salary_max']}")

#                     st.markdown("### 🔗 Apply Link")
#                     st.write(f"[View Job Posting]({row['redirected_url']})")

#                     # ======================================================
#                     # ⭐ UPDATED — Sponsorship Icon Display
#                     # ======================================================
#                     st.markdown("### 🛂 Sponsorship Eligibility")

#                     if row["is_sponsor"] is True:
#                         st.markdown(
#                             "<span style='font-size:22px;'>✅ Yes — H1B Sponsorship Available</span>",
#                             unsafe_allow_html=True
#                         )
#                     elif row["is_sponsor"] is False:
#                         st.markdown(
#                             "<span style='font-size:22px;'>❌ No — Sponsorship Not Offered</span>",
#                             unsafe_allow_html=True
#                         )
#                     else:
#                         st.markdown(
#                             "<span style='font-size:22px;'>❓ Unknown — No Sponsorship Data</span>",
#                             unsafe_allow_html=True
#                         )


# # ============================================================
# # TAB 2 — Hiring Assistant (unchanged)
# # ============================================================
# with tab2:
#     st.title("🤖 Hiring Assistant")
#     st.caption("Ask me about job descriptions, hiring messages, etc.")

#     if "messages" not in st.session_state:
#         st.session_state.messages = [{
#             "role": "assistant",
#             "content": "Hello! What role are you hiring for?"
#         }]

#     for msg in st.session_state.messages:
#         with st.chat_message(msg["role"]):
#             st.markdown(msg["content"])

#     if prompt := st.chat_input("Describe the role..."):
#         st.session_state.messages.append({"role": "user", "content": prompt})

#         with st.chat_message("assistant"):
#             with st.spinner("Searching job market data..."):
#                 context_data = retrieve_rag_context(prompt)

#                 sys_prompt = f"""
#                 Use the job market context below to write a hiring message.
#                 ---
#                 {context_data}
#                 ---
#                 User request: {prompt}
#                 """

#                 response = client.chat.completions.create(
#                     model=MODEL_NAME,
#                     messages=[{"role": "system", "content": sys_prompt}]
#                 )

#                 output = response.choices[0].message.content
#                 st.markdown(output)

#         st.session_state.messages.append({"role": "assistant", "content": output})

#######################

# import streamlit as st
# import pandas as pd
# from openai import OpenAI
# from google.cloud import bigquery
# import PyPDF2
# import docx
# import numpy as np
# from numpy.linalg import norm
# import re
# import json

# # ============================================================
# # CONFIGURATION
# # ============================================================
# PROJECT_ID = "ba882-team4-474802"
# DATASET_ID = "ba882_jobs"
# EMBED_MODEL = "text-embedding-3-small"
# MODEL_NAME = "gpt-4o-mini"

# st.set_page_config(page_title="Career Assistant Suite", layout="wide")

# # ============================================================
# # AUTH — OPENAI
# # ============================================================
# if "OPEN_AI_API" not in st.secrets:
#     st.error("❌ Missing OPEN_AI_API in secrets.toml")
# else:
#     client = OpenAI(api_key=st.secrets["OPEN_AI_API"])

# # ============================================================
# # BIGQUERY CLIENT
# # ============================================================
# client_bq = bigquery.Client(project=PROJECT_ID)


# # ============================================================
# # HELPER — Extract Resume Text
# # ============================================================
# def extract_text_from_file(uploaded_file):
#     text = ""
#     try:
#         if uploaded_file.name.endswith(".pdf"):
#             reader = PyPDF2.PdfReader(uploaded_file)
#             for page in reader.pages:
#                 extracted = page.extract_text()
#                 if extracted:
#                     text += extracted + "\n"

#         elif uploaded_file.name.endswith(".docx"):
#             doc = docx.Document(uploaded_file)
#             for para in doc.paragraphs:
#                 text += para.text + "\n"

#         else:
#             text = uploaded_file.read().decode("utf-8", errors="ignore")

#     except Exception as e:
#         st.error(f"Error reading file: {e}")

#     return text


# # ============================================================
# # SAFE JSON LIST EXTRACTOR (NEW)
# # ============================================================
# def safe_extract_keywords(text):
#     match = re.search(r"\[[^\]]*\]", text)
#     if not match:
#         return []

#     try:
#         return json.loads(match.group(0))
#     except:
#         return []


# # ============================================================
# # LLM Keyword Extraction (UPDATED)
# # ============================================================
# def extract_keywords_with_llm(resume_text, max_keywords=10):
#     prompt = f"""
#     Extract the {max_keywords} most relevant job-related keywords from the resume.

#     OUTPUT RULES:
#     - You MUST output ONLY a JSON array.
#     - No explanation, no formatting, no extra text.
#     - Example output: ["python", "sql", "tableau"]

#     Resume:
#     {resume_text}
#     """

#     response = client.chat.completions.create(
#         model=MODEL_NAME,
#         messages=[
#             {"role": "system", "content": "Output only a JSON array of keywords."},
#             {"role": "user", "content": prompt}
#         ]
#     )

#     raw = response.choices[0].message.content.strip()

#     # Direct JSON parsing
#     try:
#         keywords = json.loads(raw)
#         if isinstance(keywords, list):
#             return [k.lower() for k in keywords]
#     except:
#         pass

#     # Fallback regex parsing
#     keywords = safe_extract_keywords(raw)
#     if isinstance(keywords, list):
#         return [k.lower() for k in keywords]

#     return []


# # ============================================================
# # Embedding helpers (Resume Matcher)
# # ============================================================

# def load_job_embeddings_filtered(selected_state):
#     """
#     Loads embeddings for either:
#     - ALL states
#     - Specific state selected by user
#     """
#     if selected_state == "All States":
#         query = f"""
#             SELECT job_id, embedding_vector
#             FROM `{PROJECT_ID}.{DATASET_ID}.job_embeddings`
#         """
#     else:
#         query = f"""
#             SELECT e.job_id, e.embedding_vector
#             FROM `{PROJECT_ID}.{DATASET_ID}.job_embeddings` e
#             JOIN `{PROJECT_ID}.{DATASET_ID}.jobs` j
#             ON e.job_id = j.job_id
#             WHERE j.state = '{selected_state}'
#         """

#     return client_bq.query(query).to_dataframe()


# def fetch_job_details(job_ids):
#     formatted = ",".join([f"'{jid}'" for jid in job_ids])

#     query = f"""
#         SELECT job_id, title, company_name, city, state, description
#         FROM `{PROJECT_ID}.{DATASET_ID}.jobs`
#         WHERE job_id IN ({formatted})
#     """
#     return client_bq.query(query).to_dataframe()


# def embed_text(text):
#     emb = client.embeddings.create(
#         model=EMBED_MODEL,
#         input=text
#     )
#     return np.array(emb.data[0].embedding)


# def cosine_similarity(a, b):
#     return np.dot(a, b) / (norm(a) * norm(b))


# # ============================================================
# # RAG Context (Hiring Assistant)
# # ============================================================
# def retrieve_rag_context(query_text, top_k=3):
#     tokens = [w.lower() for w in query_text.split()]
#     keywords = [w for w in tokens if len(w) > 3]

#     if not keywords:
#         return "No meaningful terms found."

#     keyword = keywords[0]

#     query = f"""
#         SELECT title, company_name, description
#         FROM `{PROJECT_ID}.{DATASET_ID}.jobs`
#         WHERE LOWER(description) LIKE LOWER('%{keyword}%')
#            OR LOWER(title) LIKE LOWER('%{keyword}%')
#         LIMIT {int(top_k)}
#     """

#     df = client_bq.query(query).to_dataframe()

#     if df.empty:
#         return "No similar job descriptions found."

#     context_str = ""
#     for _, row in df.iterrows():
#         desc = row.get("description", "")
#         if not isinstance(desc, str):
#             desc = ""

#         context_str += f"Title: {row['title']} | Company: {row['company_name']}\n"
#         context_str += f"Description Snippet: {desc[:700]}...\n\n"

#     return context_str


# # ============================================================
# # UI: Two Tabs
# # ============================================================
# tab1, tab2 = st.tabs(["📄 Resume Matcher", "🤖 Hiring Assistant"])


# # ============================================================
# # TAB 1 — Resume Matcher (UPDATED: added State filter)
# # ============================================================
# with tab1:
#     st.title("📄 AI-Powered Resume Matcher")
#     st.caption("Upload your resume → extract skills → compute embedding → match with job embeddings.")

#     # Load distinct states for dropdown
#     states_df = client_bq.query(
#         f"SELECT DISTINCT state FROM `{PROJECT_ID}.{DATASET_ID}.jobs` WHERE state IS NOT NULL"
#     ).to_dataframe()

#     selected_state = st.selectbox("Filter by State (Optional)", ["All States"] + sorted(states_df["state"].unique()))

#     uploaded_file = st.file_uploader("Upload Resume (PDF/DOCX/TXT)", type=["pdf", "docx", "txt"])

#     if uploaded_file:
#         resume_text = extract_text_from_file(uploaded_file)
#         st.success("Resume parsed successfully!")
#         st.text_area("Extracted Resume Text", resume_text, height=200)

#         if st.button("🔍 Find Best Job Matches"):
#             with st.spinner("Processing resume and comparing embeddings..."):

#                 keywords = extract_keywords_with_llm(resume_text)

#                 st.subheader("🔑 Extracted Keywords")
#                 st.write(", ".join(keywords) if keywords else "No keywords found.")

#                 resume_embedding = embed_text(resume_text)

#                 # 🔥 Filtered job embeddings
#                 job_emb_df = load_job_embeddings_filtered(selected_state)

#                 if job_emb_df.empty:
#                     st.error("No job postings found for this state.")
#                     st.stop()

#                 similarities = []
#                 for _, row in job_emb_df.iterrows():
#                     job_vec = np.array(row["embedding_vector"])
#                     sim = cosine_similarity(resume_embedding, job_vec)
#                     similarities.append((row["job_id"], sim))

#                 similarities = sorted(similarities, key=lambda x: x[1], reverse=True)

#                 top_ids = [job_id for job_id, _ in similarities[:20]]
#                 job_details = fetch_job_details(top_ids)

#                 job_details["similarity"] = job_details["job_id"].map(dict(similarities))

#                 job_details["unique_key"] = (
#                     job_details["title"].str.lower()
#                     + " — "
#                     + job_details["company_name"].str.lower()
#                 )

#                 job_details_unique = job_details.drop_duplicates(subset="unique_key")
#                 top3 = job_details_unique.sort_values("similarity", ascending=False).head(3)

#             st.subheader("🏆 Top 3 Job Recommendations")

#             for _, row in top3.iterrows():
#                 with st.expander(
#                     f"{row['title']} — {row['company_name']}"
#                 ):
#                     st.markdown(f"**Location:** {row['city']}, {row['state']}")
#                     st.markdown(f"**Description:** {row['description'][:600]}...")


# # ============================================================
# # TAB 2 — Hiring Assistant (UNCHANGED)
# # ============================================================
# with tab2:
#     st.title("🤖 Hiring Assistant (RAG Chatbot)")
#     st.caption("Ask me to draft job descriptions or hiring messages using real job market data.")

#     if "messages" not in st.session_state:
#         st.session_state.messages = [{
#             "role": "assistant",
#             "content": "Hello! What role are you hiring for? (e.g., 'Senior Data Analyst in NYC')"
#         }]

#     for message in st.session_state.messages:
#         with st.chat_message(message["role"]):
#             st.markdown(message["content"])

#     if prompt := st.chat_input("Describe the role..."):
#         st.session_state.messages.append({"role": "user", "content": prompt})

#         with st.chat_message("user"):
#             st.markdown(prompt)

#         with st.chat_message("assistant"):
#             with st.spinner("Retrieving job market data..."):
#                 context_data = retrieve_rag_context(prompt)

#                 system_prompt = f"""
#                 You are an expert HR assistant.
#                 Use the REAL, RECENT job postings below to generate insights.

#                 --- MARKET CONTEXT ---
#                 {context_data}
#                 ----------------------

#                 USER REQUEST: {prompt}

#                 Your job: Write a polished, professional job description or hiring message.
#                 """

#                 response = client.chat.completions.create(
#                     model=MODEL_NAME,
#                     messages=[{"role": "system", "content": system_prompt}]
#                 )

#                 output = response.choices[0].message.content
#                 st.markdown(output)

#         st.session_state.messages.append({"role": "assistant", "content": output})